"""
肿瘤学全球数据到柳叶刀 - 导出服务
支持图表导出 (PNG/SVG/PDF)、论文导出 (Word/PDF)、打包下载
"""

import os
import io
import json
import zipfile
import base64
from pathlib import Path
from typing import List, Dict, Any, Optional
from datetime import datetime
import matplotlib
matplotlib.use('Agg')
import matplotlib.pyplot as plt
import matplotlib.backends.backend_svg as backend_svg


class ExportService:
    """导出服务"""
    
    def __init__(self, output_dir: str = "output"):
        self.output_dir = Path(output_dir)
        self.output_dir.mkdir(parents=True, exist_ok=True)
        
        # 创建子目录
        self.charts_dir = self.output_dir / "charts"
        self.papers_dir = self.output_dir / "papers"
        self.exports_dir = self.output_dir / "exports"
        
        for d in [self.charts_dir, self.papers_dir, self.exports_dir]:
            d.mkdir(parents=True, exist_ok=True)
    
    def export_chart_svg(
        self,
        fig: plt.Figure,
        filename: str,
        dpi: int = 300
    ) -> str:
        """
        导出图表为 SVG 格式
        
        Args:
            fig: matplotlib Figure 对象
            filename: 文件名
            dpi: 分辨率
            
        Returns:
            文件路径
        """
        filepath = self.charts_dir / f"{filename}.svg"
        fig.savefig(filepath, format='svg', dpi=dpi, bbox_inches='tight')
        return str(filepath)
    
    def export_chart_png(
        self,
        fig: plt.Figure,
        filename: str,
        dpi: int = 300,
        width: int = 1920,
        height: int = 1080
    ) -> str:
        """
        导出图表为高分辨率 PNG
        
        Args:
            fig: matplotlib Figure 对象
            filename: 文件名
            dpi: 分辨率
            width: 宽度 (像素)
            height: 高度 (像素)
            
        Returns:
            文件路径
        """
        # 设置图表大小
        fig.set_size_inches(width / dpi, height / dpi)
        
        filepath = self.charts_dir / f"{filename}.png"
        fig.savefig(filepath, format='png', dpi=dpi, bbox_inches='tight',
                   facecolor='white', edgecolor='none')
        return str(filepath)
    
    def export_chart_pdf(
        self,
        fig: plt.Figure,
        filename: str
    ) -> str:
        """
        导出图表为 PDF 格式
        
        Args:
            fig: matplotlib Figure 对象
            filename: 文件名
            
        Returns:
            文件路径
        """
        filepath = self.charts_dir / f"{filename}.pdf"
        fig.savefig(filepath, format='pdf', bbox_inches='tight')
        return str(filepath)
    
    def export_paper_word(
        self,
        title: str,
        sections: Dict[str, Any],
        filename: str
    ) -> str:
        """
        导出论文为 Word 格式
        
        Args:
            title: 论文标题
            sections: 论文章节内容
            filename: 文件名
            
        Returns:
            文件路径
        """
        try:
            from docx import Document
            from docx.shared import Inches, Pt, RGBColor
            from docx.enum.text import WD_ALIGN_PARAGRAPH
            
            doc = Document()
            
            # 设置标题
            title_para = doc.add_heading(title, 0)
            title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            # 添加作者和日期
            doc.add_paragraph(f"Authors: AI Research Assistant")
            doc.add_paragraph(f"Date: {datetime.now().strftime('%Y-%m-%d')}")
            doc.add_paragraph("---")
            
            # Summary
            doc.add_heading("Summary", level=1)
            if isinstance(sections.get("summary"), dict):
                for key, value in sections["summary"].items():
                    doc.add_paragraph(f"{key.title()}: {value}")
            else:
                doc.add_paragraph(str(sections.get("summary", "")))
            
            # Introduction
            doc.add_heading("Introduction", level=1)
            doc.add_paragraph(str(sections.get("introduction", "")))
            
            # Methods
            doc.add_heading("Methods", level=1)
            doc.add_paragraph(str(sections.get("methods", "")))
            
            # Results
            doc.add_heading("Results", level=1)
            doc.add_paragraph(str(sections.get("results", "")))
            
            # Discussion
            doc.add_heading("Discussion", level=1)
            doc.add_paragraph(str(sections.get("discussion", "")))
            
            # 保存
            filepath = self.papers_dir / f"{filename}.docx"
            doc.save(str(filepath))
            
            return str(filepath)
            
        except ImportError:
            # 如果没有 python-docx，使用纯文本
            return self.export_paper_text(title, sections, filename)
    
    def export_paper_text(
        self,
        title: str,
        sections: Dict[str, Any],
        filename: str
    ) -> str:
        """
        导出论文为纯文本格式
        
        Args:
            title: 论文标题
            sections: 论文章节内容
            filename: 文件名
            
        Returns:
            文件路径
        """
        filepath = self.papers_dir / f"{filename}.txt"
        
        with open(filepath, 'w', encoding='utf-8') as f:
            f.write(f"{title}\n")
            f.write("=" * len(title) + "\n\n")
            f.write(f"Authors: AI Research Assistant\n")
            f.write(f"Date: {datetime.now().strftime('%Y-%m-%d')}\n")
            f.write("-" * 50 + "\n\n")
            
            for section_name, content in sections.items():
                f.write(f"\n{section_name.upper()}\n")
                f.write("-" * 20 + "\n")
                
                if isinstance(content, dict):
                    for key, value in content.items():
                        f.write(f"\n{key.title()}:\n{value}\n")
                else:
                    f.write(f"{content}\n")
        
        return str(filepath)
    
    def export_paper_html(
        self,
        title: str,
        sections: Dict[str, Any],
        filename: str
    ) -> str:
        """
        导出论文为 HTML 格式
        
        Args:
            title: 论文标题
            sections: 论文章节内容
            filename: 文件名
            
        Returns:
            文件路径
        """
        filepath = self.papers_dir / f"{filename}.html"
        
        html = f"""<!DOCTYPE html>
<html lang="en">
<head>
    <meta charset="UTF-8">
    <meta name="viewport" content="width=device-width, initial-scale=1.0">
    <title>{title}</title>
    <style>
        body {{
            font-family: 'Times New Roman', serif;
            max-width: 800px;
            margin: 0 auto;
            padding: 40px;
            line-height: 1.8;
            color: #333;
        }}
        h1 {{
            color: #A51C30;
            border-bottom: 2px solid #A51C30;
            padding-bottom: 10px;
        }}
        h2 {{
            color: #2166AC;
            margin-top: 30px;
        }}
        .meta {{
            color: #666;
            font-style: italic;
            margin-bottom: 30px;
        }}
        hr {{
            border: none;
            border-top: 1px solid #ddd;
            margin: 30px 0;
        }}
        @media print {{
            body {{
                max-width: 100%;
            }}
        }}
    </style>
</head>
<body>
    <h1>{title}</h1>
    <div class="meta">
        <p><strong>Authors:</strong> AI Research Assistant</p>
        <p><strong>Date:</strong> {datetime.now().strftime('%Y-%m-%d')}</p>
    </div>
    <hr>
"""
        
        for section_name, content in sections.items():
            html += f"<h2>{section_name.title()}</h2>\n"
            if isinstance(content, dict):
                for key, value in content.items():
                    html += f"<p><strong>{key.title()}:</strong> {value}</p>\n"
            else:
                # 简单的段落处理
                paragraphs = str(content).split('\n')
                for para in paragraphs:
                    if para.strip():
                        html += f"<p>{para}</p>\n"
            html += "<hr>\n"
        
        html += """
</body>
</html>"""
        
        with open(filepath, 'w', encoding='utf-8') as f:
            f.write(html)
        
        return str(filepath)
    
    def create_download_package(
        self,
        project_id: str,
        paper_files: List[str] = None,
        chart_files: List[str] = None,
        data_files: List[str] = None,
        include_analysis: bool = True
    ) -> str:
        """
        创建下载包 (ZIP)
        
        Args:
            project_id: 项目 ID
            paper_files: 论文文件路径列表
            chart_files: 图表文件路径列表
            data_files: 数据文件路径列表
            include_analysis: 是否包含分析结果
            
        Returns:
            ZIP 文件路径
        """
        timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
        zip_filename = f"{project_id}_{timestamp}.zip"
        zip_filepath = self.exports_dir / zip_filename
        
        with zipfile.ZipFile(zip_filepath, 'w', zipfile.ZIP_DEFLATED) as zipf:
            # 添加论文文件
            if paper_files:
                for filepath in paper_files:
                    if os.path.exists(filepath):
                        arcname = f"papers/{os.path.basename(filepath)}"
                        zipf.write(filepath, arcname)
            
            # 添加图表文件
            if chart_files:
                for filepath in chart_files:
                    if os.path.exists(filepath):
                        arcname = f"charts/{os.path.basename(filepath)}"
                        zipf.write(filepath, arcname)
            
            # 添加数据文件
            if data_files:
                for filepath in data_files:
                    if os.path.exists(filepath):
                        arcname = f"data/{os.path.basename(filepath)}"
                        zipf.write(filepath, arcname)
            
            # 添加 README
            readme_content = f"""# Cancer Epidemiology Research To Lancet
# 肿瘤学全球数据到柳叶刀

Project ID: {project_id}
Export Date: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}

## Contents

"""
            if paper_files:
                readme_content += "## Papers\n"
                for f in paper_files:
                    readme_content += f"- {os.path.basename(f)}\n"
            
            if chart_files:
                readme_content += "\n## Charts\n"
                for f in chart_files:
                    readme_content += f"- {os.path.basename(f)}\n"
            
            readme_content += """
## Citation

If you use this data in your research, please cite:
Cancer Epidemiology Research To Lancet Platform
"""
            
            zipf.writestr("README.md", readme_content)
        
        return str(zip_filepath)
    
    def get_export_summary(
        self,
        project_id: str
    ) -> Dict[str, Any]:
        """
        获取项目导出文件摘要
        
        Args:
            project_id: 项目 ID
            
        Returns:
            导出文件摘要
        """
        summary = {
            "project_id": project_id,
            "charts": [],
            "papers": [],
            "exports": []
        }
        
        # 扫描图表文件
        for f in self.charts_dir.glob(f"{project_id}*"):
            summary["charts"].append({
                "filename": f.name,
                "size": f.stat().st_size,
                "path": str(f)
            })
        
        # 扫描论文文件
        for f in self.papers_dir.glob(f"{project_id}*"):
            summary["papers"].append({
                "filename": f.name,
                "size": f.stat().st_size,
                "path": str(f)
            })
        
        # 扫描导出包
        for f in self.exports_dir.glob(f"{project_id}*"):
            summary["exports"].append({
                "filename": f.name,
                "size": f.stat().st_size,
                "path": str(f)
            })
        
        return summary


# 创建全局实例
export_service = ExportService()
